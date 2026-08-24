from io import BytesIO
from zipfile import is_zipfile

from conftest import register_and_login
from docx import Document

from studyai.db import get_db
from studyai.jobs import COMPLETED, UPLOADING, create_job
from studyai.services.exports import organize_transcript, parse_study_blocks

EXPLANATION = """# المفهوم الأساسي

هذا شرح عربي منظم يوضح الفكرة بصورة مباشرة.

- النقطة الأولى
- النقطة الثانية

1. الخطوة الأولى
2. الخطوة الثانية

[مثال توضيحي مضاف] إذا زادت السرعة، تقل مدة الرحلة.
"""


def _completed_job(app, *, explanation: str | None = EXPLANATION) -> str:
    with app.app_context():
        user_id = get_db().execute(
            "SELECT id FROM users WHERE username='__public__'"
        ).fetchone()[0]
        job_id = create_job(user_id, "محاضرة الفيزياء.mp3", status=UPLOADING)
        database = get_db()
        database.execute(
            """UPDATE processing_jobs SET status = ?, current_stage = ?, progress = 100,
               transcript = ?, summary = ?, questions = ?, explanation = ? WHERE id = ?""",
            (
                COMPLETED,
                COMPLETED,
                "هذا هو التفريغ الكامل للمحاضرة.",
                "# الملخص\n\n- الفكرة الأساسية",
                "1. ما الفكرة الأساسية؟\n\nالإجابة: الفكرة الأساسية.",
                explanation,
                job_id,
            ),
        )
        database.commit()
        return job_id


def test_explanation_exports_are_real_downloadable_formats(app, client):
    register_and_login(client)
    job_id = _completed_job(app)

    markdown = client.get(f"/api/jobs/{job_id}/export/explanation.md")
    assert markdown.status_code == 200
    assert markdown.data.startswith(b"\xef\xbb\xbf# ")
    assert "attachment" in markdown.headers["Content-Disposition"]
    assert markdown.headers["Cache-Control"] == "no-store"

    word = client.get(f"/api/jobs/{job_id}/export/explanation.docx")
    assert word.status_code == 200
    assert is_zipfile(BytesIO(word.data))
    document = Document(BytesIO(word.data))
    assert any("الشرح والأمثلة التوضيحية" in p.text for p in document.paragraphs)
    assert any("النقطة الأولى" in p.text for p in document.paragraphs)

    pdf = client.get(f"/api/jobs/{job_id}/export/explanation.pdf")
    assert pdf.status_code == 200
    assert pdf.data.startswith(b"%PDF-")
    assert pdf.data.rstrip().endswith(b"%%EOF")
    assert b"/Type /Page" in pdf.data


def test_transcript_summary_and_questions_can_each_be_exported(app, client):
    register_and_login(client)
    job_id = _completed_job(app)

    cases = {
        "transcript": "التفريغ الكامل",
        "summary": "ملخص المحاضرة",
        "questions": "أسئلة المراجعة وإجاباتها",
    }
    for result_type, title in cases.items():
        markdown = client.get(f"/api/jobs/{job_id}/export/{result_type}.md")
        assert markdown.status_code == 200
        assert title.encode("utf-8") in markdown.data
        if result_type == "transcript":
            assert "## القسم 01".encode() in markdown.data

        word = client.get(f"/api/jobs/{job_id}/export/{result_type}.docx")
        assert word.status_code == 200
        assert is_zipfile(BytesIO(word.data))

        pdf = client.get(f"/api/jobs/{job_id}/export/{result_type}.pdf")
        assert pdf.status_code == 200
        assert pdf.data.startswith(b"%PDF-")


def test_explanation_export_validates_owner_status_content_and_format(app, client):
    register_and_login(client)
    empty_job_id = _completed_job(app, explanation=None)
    assert client.get(f"/api/jobs/{empty_job_id}/export/explanation.pdf").status_code == 409
    assert client.get(f"/api/jobs/{empty_job_id}/export/explanation.txt").status_code == 409
    assert client.get(f"/api/jobs/{empty_job_id}/export/unknown.md").status_code == 404

    with client.session_transaction() as user_session:
        user_session.clear()
    assert client.get(f"/api/jobs/{empty_job_id}/export/explanation.md").status_code == 404


def test_study_block_parser_preserves_semantic_structure():
    blocks = parse_study_blocks(EXPLANATION)
    assert [block.kind for block in blocks] == [
        "heading",
        "paragraph",
        "bullet",
        "bullet",
        "numbered",
        "numbered",
        "example",
    ]
    assert blocks[-1].text.startswith("إذا زادت السرعة")


def test_transcript_organization_preserves_every_original_word_and_grounds_titles():
    transcript = " ".join(f"كلمة{i}" for i in range(905))
    sections = organize_transcript(transcript)
    assert len(sections) == 3
    assert " ".join(section.text for section in sections) == transcript
    for section in sections:
        excerpt = section.title.split("—", 1)[1].strip().removesuffix("…")
        assert section.text.startswith(excerpt)


def test_result_api_keeps_raw_transcript_and_adds_safe_display_sections(app, client):
    register_and_login(client)
    job_id = _completed_job(app)
    payload = client.get(f"/api/jobs/{job_id}/result").get_json()
    assert payload["transcript"] == "هذا هو التفريغ الكامل للمحاضرة."
    assert " ".join(item["text"] for item in payload["transcript_sections"]) == payload[
        "transcript"
    ]
