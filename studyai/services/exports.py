"""Create polished exports for every completed lecture result."""

from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path

import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

GREEN = "1F6B4F"
GREEN_DARK = "173C2F"
GREEN_LIGHT = "E8F2E7"
INK = "18231D"
MUTED = "66716B"
FONT_PATH = (
    Path(__file__).resolve().parents[1] / "static" / "fonts" / "NotoSansArabic-Regular.ttf"
)

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
_BULLET_RE = re.compile(r"^[-*•]\s+(.+)$")
_NUMBER_RE = re.compile(r"^(\d+)[.)،-]\s+(.+)$")
_EXAMPLE_RE = re.compile(r"^\s*(?:\[)?مثال(?:\s+توضيحي)?(?:\s+مضاف)?(?:\])?\s*[:：-]?\s*(.*)$")
_INLINE_MARKDOWN_RE = re.compile(r"(?:\*\*|__|`|~~)")
_LINK_RE = re.compile(r"\[([^]]+)]\(([^)]+)\)")


@dataclass(frozen=True)
class StudyBlock:
    kind: str
    text: str
    level: int = 0
    marker: str = ""


def build_markdown(
    content_text: str,
    source_name: str,
    title: str = "الشرح والأمثلة التوضيحية",
) -> bytes:
    """Return editable UTF-8 Markdown with a small, useful document header."""
    source = source_name.strip() or "محاضرة"
    body = content_text.strip()
    generated = datetime.now(UTC).strftime("%Y-%m-%d")
    content = (
        f"# {title}\n\n"
        f"> **المصدر:** {source}  \n"
        f"> **تاريخ الإنشاء:** {generated}  \n"
        "> **أُنشئ بواسطة:** StudyAI\n\n"
        "---\n\n"
        f"{body}\n"
    )
    return b"\xef\xbb\xbf" + content.replace("\r\n", "\n").encode("utf-8")


def build_docx(
    content_text: str,
    source_name: str,
    title_text: str = "الشرح والأمثلة التوضيحية",
) -> bytes:
    """Return a real RTL Word document using a compact reference-guide layout."""
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    _configure_docx_styles(document)
    _add_docx_header_footer(section)

    kicker = document.add_paragraph()
    _format_docx_paragraph(kicker, before=0, after=4)
    kicker.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = kicker.add_run("دليل دراسي منظم")
    _format_docx_run(run, size=10, color=GREEN, bold=True)

    title = document.add_paragraph(style="Title")
    _format_docx_paragraph(title, before=0, after=8, line_spacing=1.0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = title.add_run(title_text)
    _format_docx_run(run, size=28, color=GREEN_DARK, bold=True)

    metadata = document.add_paragraph()
    _format_docx_paragraph(metadata, before=0, after=18)
    metadata.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    source = source_name.strip() or "محاضرة"
    run = metadata.add_run(
        f"المصدر: {source}  |  أُنشئ بواسطة StudyAI  |  "
        f"{datetime.now(UTC).strftime('%Y-%m-%d')}"
    )
    _format_docx_run(run, size=9.5, color=MUTED)

    previous_kind = ""
    numbering_id: int | None = None
    for block in parse_study_blocks(content_text):
        if block.kind == "heading":
            paragraph = document.add_paragraph(style=f"Heading {min(block.level, 3)}")
            run = paragraph.add_run(block.text)
            _format_docx_run(
                run,
                size={1: 16, 2: 13, 3: 12}.get(block.level, 12),
                color=GREEN if block.level < 3 else GREEN_DARK,
                bold=True,
            )
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(block.text)
            _format_docx_run(run)
        elif block.kind == "numbered":
            if previous_kind != "numbered":
                numbering_id = _create_numbered_list(document, int(block.marker or "1"))
            paragraph = document.add_paragraph(style="StudyAI Numbered")
            _apply_numbering(paragraph, numbering_id)
            run = paragraph.add_run(block.text)
            _format_docx_run(run)
        elif block.kind == "example":
            paragraph = document.add_paragraph(style="StudyAI Example")
            run = paragraph.add_run("مثال توضيحي\n")
            _format_docx_run(run, size=10, color=GREEN, bold=True)
            run = paragraph.add_run(block.text)
            _format_docx_run(run)
            _shade_paragraph(paragraph, GREEN_LIGHT)
        else:
            paragraph = document.add_paragraph(style="Normal")
            run = paragraph.add_run(block.text)
            _format_docx_run(run)
        _set_paragraph_rtl(paragraph)
        previous_kind = block.kind

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_pdf(
    content_text: str,
    source_name: str,
    title_text: str = "الشرح والأمثلة التوضيحية",
) -> bytes:
    """Return a paginated Arabic PDF with shaped, line-wrapped RTL text."""
    _register_pdf_font()
    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=LETTER, pageCompression=1)
    pdf.setTitle(f"StudyAI - {title_text}")
    pdf.setAuthor("StudyAI")
    renderer = _PdfStudyGuide(pdf)
    renderer.start_page()
    renderer.draw_title(source_name, title_text)
    for block in parse_study_blocks(content_text):
        renderer.draw_block(block)
    renderer.finish()
    return output.getvalue()


def parse_study_blocks(text: str) -> list[StudyBlock]:
    """Parse the common Markdown subset produced by the AI into semantic blocks."""
    blocks: list[StudyBlock] = []
    paragraph_lines: list[str] = []
    in_code_fence = False

    def flush_paragraph() -> None:
        if paragraph_lines:
            value = _clean_inline(" ".join(paragraph_lines))
            if value:
                blocks.append(StudyBlock("paragraph", value))
            paragraph_lines.clear()

    for raw_line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.strip()
        if line.startswith("```"):
            flush_paragraph()
            in_code_fence = not in_code_fence
            continue
        if not line:
            flush_paragraph()
            continue
        if in_code_fence:
            paragraph_lines.append(line)
            continue
        heading = _HEADING_RE.match(line)
        bullet = _BULLET_RE.match(line)
        numbered = _NUMBER_RE.match(line)
        example = _EXAMPLE_RE.match(_clean_inline(line))
        if heading:
            flush_paragraph()
            blocks.append(
                StudyBlock("heading", _clean_inline(heading.group(2)), len(heading.group(1)))
            )
        elif bullet:
            flush_paragraph()
            blocks.append(StudyBlock("bullet", _clean_inline(bullet.group(1))))
        elif numbered:
            flush_paragraph()
            blocks.append(
                StudyBlock("numbered", _clean_inline(numbered.group(2)), marker=numbered.group(1))
            )
        elif example and example.group(1):
            flush_paragraph()
            blocks.append(StudyBlock("example", _clean_inline(example.group(1))))
        else:
            paragraph_lines.append(line)
    flush_paragraph()
    return blocks


def _clean_inline(value: str) -> str:
    value = _LINK_RE.sub(r"\1 (\2)", value)
    return _INLINE_MARKDOWN_RE.sub("", value).strip()


def _configure_docx_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    _set_style_rtl(normal)

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(28)
    title.font.color.rgb = RGBColor.from_string(GREEN_DARK)
    title.font.bold = True
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    _set_style_rtl(title)
    title_props = title.element.get_or_add_pPr()
    title_border = title_props.find(qn("w:pBdr"))
    if title_border is not None:
        title_props.remove(title_border)

    heading_tokens = {
        "Heading 1": (16, GREEN, 18, 10),
        "Heading 2": (13, GREEN, 14, 7),
        "Heading 3": (12, GREEN_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        _set_style_rtl(style)

    for name in ("List Bullet",):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.right_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        _set_style_rtl(style)

    if "StudyAI Numbered" not in styles:
        numbered = styles.add_style("StudyAI Numbered", WD_STYLE_TYPE.PARAGRAPH)
    else:
        numbered = styles["StudyAI Numbered"]
    numbered.base_style = normal
    numbered.font.name = "Arial"
    numbered.font.size = Pt(11)
    numbered.paragraph_format.space_before = Pt(0)
    numbered.paragraph_format.space_after = Pt(4)
    numbered.paragraph_format.line_spacing = 1.25
    numbered.paragraph_format.right_indent = Inches(0.375)
    numbered.paragraph_format.first_line_indent = Inches(-0.188)
    _set_style_rtl(numbered)

    if "StudyAI Example" not in styles:
        example = styles.add_style("StudyAI Example", WD_STYLE_TYPE.PARAGRAPH)
    else:
        example = styles["StudyAI Example"]
    example.base_style = normal
    example.font.name = "Arial"
    example.font.size = Pt(11)
    example.paragraph_format.space_before = Pt(6)
    example.paragraph_format.space_after = Pt(8)
    example.paragraph_format.line_spacing = 1.25
    example.paragraph_format.left_indent = Inches(0.12)
    example.paragraph_format.right_indent = Inches(0.12)
    _set_style_rtl(example)


def _format_docx_paragraph(
    paragraph, *, before: float = 0, after: float = 6, line_spacing: float = 1.25
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    _set_paragraph_rtl(paragraph)


def _format_docx_run(
    run, *, size: float = 11, color: str = INK, bold: bool | None = None
) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    run_props = run._element.get_or_add_rPr()
    fonts = run_props.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_props.insert(0, fonts)
    for key in ("ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{key}"), "Arial")
    if run_props.find(qn("w:rtl")) is None:
        run_props.append(OxmlElement("w:rtl"))


def _set_paragraph_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    props = paragraph._element.get_or_add_pPr()
    if props.find(qn("w:bidi")) is None:
        props.append(OxmlElement("w:bidi"))


def _set_style_rtl(style) -> None:
    props = style.element.get_or_add_pPr()
    if props.find(qn("w:bidi")) is None:
        props.append(OxmlElement("w:bidi"))


def _shade_paragraph(paragraph, fill: str) -> None:
    props = paragraph._element.get_or_add_pPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), fill)
    spacing = props.find(qn("w:spacing"))
    if spacing is not None:
        spacing.set(qn("w:before"), "120")
        spacing.set(qn("w:after"), "160")


def _add_docx_header_footer(section) -> None:
    header = section.header.paragraphs[0]
    _format_docx_paragraph(header, before=0, after=0, line_spacing=1.0)
    run = header.add_run("StudyAI  |  دليل دراسي")
    _format_docx_run(run, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    _format_docx_paragraph(footer, before=0, after=0, line_spacing=1.0)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("StudyAI  |  صفحة ")
    _format_docx_run(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def _create_numbered_list(document: Document, start: int) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))
    ]
    number_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    number_id = max(number_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "decimal"), ("w:lvlText", "%1.")):
        element = OxmlElement(tag)
        element.set(qn("w:val"), value)
        level.append(element)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "right")
    level.append(justification)
    paragraph_props = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:right"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_props.append(indent)
    level.append(paragraph_props)
    run_props = OxmlElement("w:rPr")
    rtl = OxmlElement("w:rtl")
    run_props.append(rtl)
    level.append(run_props)
    abstract.append(level)
    numbering.insert(0, abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(number_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(max(1, start)))
    override.append(start_override)
    number.append(override)
    numbering.append(number)
    return number_id


def _apply_numbering(paragraph, number_id: int | None) -> None:
    if number_id is None:
        return
    props = paragraph._element.get_or_add_pPr()
    number_props = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(number_id))
    number_props.append(level)
    number_props.append(number)
    props.append(number_props)


def _register_pdf_font() -> None:
    if "StudyAICairo" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("StudyAICairo", str(FONT_PATH)))


def _rtl(value: str) -> str:
    return get_display(arabic_reshaper.reshape(value), base_dir="R")


class _PdfStudyGuide:
    font_name = "StudyAICairo"
    page_width, page_height = LETTER
    right = page_width - 72
    left = 72
    content_width = page_width - 144

    def __init__(self, pdf: canvas.Canvas) -> None:
        self.pdf = pdf
        self.page_number = 0
        self.y = 0.0

    def start_page(self) -> None:
        if self.page_number:
            self.pdf.showPage()
        self.page_number += 1
        self.pdf.setFillColor(HexColor(f"#{MUTED}"))
        self.pdf.setFont(self.font_name, 8.5)
        self.pdf.drawRightString(self.right, self.page_height - 39, _rtl("StudyAI | دليل دراسي"))
        self.pdf.drawCentredString(
            self.page_width / 2, 34, _rtl(f"StudyAI | صفحة {self.page_number}")
        )
        self.y = self.page_height - 72

    def finish(self) -> None:
        self.pdf.save()

    def draw_title(self, source_name: str, title_text: str) -> None:
        self.pdf.setFillColor(HexColor(f"#{GREEN}"))
        self.pdf.setFont(self.font_name, 10)
        self.pdf.drawRightString(self.right, self.y, _rtl("دليل دراسي منظم"))
        self.y -= 26
        self.pdf.setFillColor(HexColor(f"#{GREEN_DARK}"))
        self.pdf.setFont(self.font_name, 22)
        self.pdf.drawRightString(self.right, self.y, _rtl(title_text))
        self.y -= 31
        source = source_name.strip() or "محاضرة"
        metadata = (
            f"المصدر: {source} | أُنشئ بواسطة StudyAI | "
            f"{datetime.now(UTC).strftime('%Y-%m-%d')}"
        )
        self._draw_text(metadata, size=8.5, color=MUTED, leading=13, after=18)

    def draw_block(self, block: StudyBlock) -> None:
        if block.kind == "heading":
            size = {1: 16, 2: 13, 3: 11.5}.get(block.level, 11.5)
            before = {1: 13, 2: 10, 3: 8}.get(block.level, 8)
            self._draw_text(
                block.text,
                size=size,
                color=GREEN if block.level < 3 else GREEN_DARK,
                leading=size * 1.45,
                before=before,
                after=6,
                keep_next=True,
            )
        elif block.kind == "bullet":
            self._draw_list_item(block.text, "•")
        elif block.kind == "numbered":
            self._draw_list_item(block.text, f"{block.marker}.")
        elif block.kind == "example":
            self._draw_example(block.text)
        else:
            self._draw_text(block.text, size=10.5, color=INK, leading=17, after=7)

    def _draw_text(
        self,
        text: str,
        *,
        size: float,
        color: str,
        leading: float,
        before: float = 0,
        after: float = 0,
        keep_next: bool = False,
        right: float | None = None,
        width: float | None = None,
    ) -> None:
        target_right = right if right is not None else self.right
        max_width = width if width is not None else self.content_width
        lines = self._wrap_lines(text, size, max_width)
        needed = before + len(lines) * leading + after + (leading if keep_next else 0)
        self._ensure_space(needed)
        self.y -= before
        self.pdf.setFillColor(HexColor(f"#{color}"))
        self.pdf.setFont(self.font_name, size)
        for line in lines:
            self.pdf.drawRightString(target_right, self.y, _rtl(line))
            self.y -= leading
        self.y -= after

    def _draw_list_item(self, text: str, marker: str) -> None:
        marker_width = 22
        lines = self._wrap_lines(text, 10.5, self.content_width - marker_width)
        leading = 16.5
        self._ensure_space(len(lines) * leading + 5)
        start_y = self.y
        self.pdf.setFillColor(HexColor(f"#{GREEN}"))
        self.pdf.setFont(self.font_name, 10.5)
        self.pdf.drawRightString(self.right, start_y, _rtl(marker))
        self.pdf.setFillColor(HexColor(f"#{INK}"))
        for line in lines:
            self.pdf.drawRightString(self.right - marker_width, self.y, _rtl(line))
            self.y -= leading
        self.y -= 5

    def _draw_example(self, text: str) -> None:
        lines = self._wrap_lines(text, 10.25, self.content_width - 28)
        leading = 16.5
        label_height = 18
        full_box_height = label_height + len(lines) * leading + 18
        full_page_capacity = self.page_height - 72 - 58
        if full_box_height <= full_page_capacity and self.y - full_box_height < 58:
            self.start_page()
        while lines:
            available = self.y - 58
            max_lines = max(1, int((available - label_height - 18) // leading))
            chunk, lines = lines[:max_lines], lines[max_lines:]
            box_height = label_height + len(chunk) * leading + 18
            self._ensure_space(box_height + 8)
            top = self.y
            self.pdf.setFillColor(HexColor(f"#{GREEN_LIGHT}"))
            self.pdf.roundRect(
                self.left, top - box_height + 5, self.content_width, box_height, 8, fill=1, stroke=0
            )
            self.pdf.setFillColor(HexColor(f"#{GREEN}"))
            self.pdf.setFont(self.font_name, 9)
            self.pdf.drawRightString(self.right - 14, top - 12, _rtl("مثال توضيحي"))
            self.pdf.setFillColor(HexColor(f"#{INK}"))
            self.pdf.setFont(self.font_name, 10.25)
            y = top - label_height - 8
            for line in chunk:
                self.pdf.drawRightString(self.right - 14, y, _rtl(line))
                y -= leading
            self.y = top - box_height - 5
            if lines:
                self.start_page()

    def _wrap_lines(self, text: str, size: float, max_width: float) -> list[str]:
        words = text.split()
        if not words:
            return [""]
        lines: list[str] = []
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            width = pdfmetrics.stringWidth(_rtl(candidate), self.font_name, size)
            if width <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
        return lines

    def _ensure_space(self, needed: float) -> None:
        if self.y - needed < 58:
            self.start_page()
